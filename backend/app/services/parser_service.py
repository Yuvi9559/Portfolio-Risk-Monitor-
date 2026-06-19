import csv
import io
import re
import logging
from typing import List, Dict, Any, Optional

import pypdf
import docx

logger = logging.getLogger(__name__)

TICKER_REGEX = re.compile(r'^([A-Z]{2,6}|[A-Z]{2,6}[.-][A-Z]{1,4})$')
# Regex to match numeric values (including floats and comma-separated integers)
NUMERIC_REGEX = re.compile(r'^\$?([\d,]+(?:\.\d+)?)$')

STOP_WORDS = {
    "AND", "OR", "BUT", "THE", "FOR", "WITH", "BY", "OF", "TO", "IN", "ON", "AT", "AN", "A", 
    "IS", "ARE", "WAS", "WERE", "BE", "BEING", "BEEN", "HAVE", "HAS", "HAD", "DO", "DOES", "DID", 
    "AS", "IF", "THIS", "THAT", "THESE", "THOSE", "EACH", "ALL", "ANY", "BOTH", "SOME", "MANY", 
    "FEW", "MORE", "MOST", "OTHER", "SUCH", "ONLY", "OWN", "VERY", "CAN", "WILL", "JUST", "SHOULD", 
    "COULD", "WOULD", "SHARES", "COST", "PRICE", "TOTAL", "VALUE", "USD", "EUR", "DATE", "ANN", 
    "VOL", "BETA", "SHARPE", "MAX", "RISK", "LEVEL", "PORT", "MKT", "CAP", "SECTOR", "DIV", "YIELD",
    "HIGH", "LOW", "VERY", "WE", "ALSO", "WHICH", "WHO", "HOW", "WHY", "WHERE", "WHEN", "OUR", 
    "YOU", "THEY", "HE", "SHE", "HIM", "HER", "THEM", "US", "ME", "MY", "THEIR", "YOUR", "ABOUT", 
    "ABOVE", "AFTER", "BEFORE", "ONCE", "THAN", "THEN", "BECAUSE", "SINCE", "UNTIL", "WHILE", 
    "EITHER", "NEITHER", "NOR", "EVERY", "SAME", "SO", "TOO", "MAY", "MIGHT", "MUST", "SHALL",
    "ITS", "IT", "OUT", "INTO", "OVER", "UNDER", "AGAIN", "FURTHER", "ONCE", "HERE", "THERE"
}

def clean_numeric(val_str: str) -> Optional[float]:
    """Clean dollar signs, commas, and parse as float."""
    if not val_str:
        return None
    val_clean = val_str.replace('$', '').replace(',', '').strip()
    try:
        return float(val_clean)
    except ValueError:
        return None

def parse_csv(text: str) -> List[Dict[str, Any]]:
    """Parse CSV text, heuristically mapping columns."""
    holdings = []
    lines = text.splitlines()
    if not lines:
        return holdings

    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return holdings

    # Find headers row (usually first row, let's scan first 3 rows just in case)
    headers = []
    header_row_idx = 0
    for idx, row in enumerate(rows[:3]):
        row_lower = [c.lower().strip() for c in row]
        if any(h in c for c in row_lower for h in ['symbol', 'ticker', 'shares', 'qty', 'quantity']):
            headers = row_lower
            header_row_idx = idx
            break
    
    if not headers:
        # Default header mapping by index if headers not explicitly found
        headers = ['symbol', 'shares', 'avg_cost']

    # Map column indexes
    sym_idx = next((i for i, h in enumerate(headers) if 'symbol' in h or 'ticker' in h or 'asset' in h), 0)
    shares_idx = next((i for i, h in enumerate(headers) if 'shares' in h or 'qty' in h or 'quantity' in h), 1)
    cost_idx = next((i for i, h in enumerate(headers) if 'cost' in h or 'price' in h or 'avg' in h), -1)

    for row in rows[header_row_idx + 1:]:
        if not row or len(row) <= max(sym_idx, shares_idx):
            continue
        
        symbol = row[sym_idx].strip().upper()
        shares_val = clean_numeric(row[shares_idx])
        avg_cost_val = clean_numeric(row[cost_idx]) if cost_idx != -1 and cost_idx < len(row) else None

        if not symbol or shares_val is None or shares_val <= 0:
            continue

        asset_type = "crypto" if (symbol.endswith("-USD") or symbol in ["BTC", "ETH", "SOL", "XRP"]) else "stock"
        holdings.append({
            "symbol": symbol,
            "shares": shares_val,
            "avg_cost": avg_cost_val,
            "asset_type": asset_type
        })
    
    return holdings

def extract_holdings_from_text(text: str) -> List[Dict[str, Any]]:
    """Scan raw text lines and extract holdings matching symbol + quantity + cost patterns."""
    holdings = []
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Look for words in line (split by space/semicolon/pipe, strip commas, colons, and dots from ends)
        tokens = [t.strip(' ,;:|.') for t in re.split(r'[\s;|]+', line) if t.strip(' ,;:|.')]
        if len(tokens) < 2:
            continue

        # Look for a ticker symbol token
        symbol = None
        symbol_idx = -1
        for idx, token in enumerate(tokens):
            clean_token = token.upper().replace(':', '')
            if TICKER_REGEX.match(clean_token) and not clean_token in STOP_WORDS:
                symbol = clean_token
                symbol_idx = idx
                break
        
        if symbol is None or symbol_idx == -1:
            continue

        # Collect all numeric values with their indices
        num_tokens = []
        for idx, token in enumerate(tokens):
            if idx == symbol_idx:
                continue
            match = NUMERIC_REGEX.match(token)
            if match:
                val = clean_numeric(match.group(1))
                if val is not None and val > 0:
                    num_tokens.append((idx, val))

        pre_nums = [(idx, val) for idx, val in num_tokens if idx < symbol_idx]
        post_nums = [(idx, val) for idx, val in num_tokens if idx > symbol_idx]

        shares_val = None
        avg_cost_val = None

        if pre_nums and post_nums:
            # Case 1: Numbers exist both before and after the symbol
            shares_val = pre_nums[-1][1]
            avg_cost_val = post_nums[0][1]
        elif post_nums:
            # Case 2: Numbers only exist after the symbol
            shares_val = post_nums[0][1]
            if len(post_nums) >= 2:
                avg_cost_val = post_nums[1][1]
        elif pre_nums:
            # Case 3: Numbers only exist before the symbol
            if len(pre_nums) >= 2:
                shares_val = pre_nums[-2][1]
                avg_cost_val = pre_nums[-1][1]
            else:
                shares_val = pre_nums[0][1]

        if shares_val is None:
            continue

        # Auto-correct dot tickers to dashes (e.g. BRK.B -> BRK-B)
        symbol = symbol.replace('.', '-')
        
        # Auto-correct cryptocurrency symbols to append -USD
        KNOWN_CRYPTOS = {"BTC", "ETH", "SOL", "BNB", "LINK", "XRP", "ADA", "DOT"}
        if symbol in KNOWN_CRYPTOS:
            symbol = f"{symbol}-USD"

        asset_type = "crypto" if (symbol.endswith("-USD") or symbol in ["BTC", "ETH", "SOL", "XRP"]) else "stock"
        holdings.append({
            "symbol": symbol,
            "shares": shares_val,
            "avg_cost": avg_cost_val,
            "asset_type": asset_type
        })

    # Remove duplicates by merging shares if the same symbol is extracted multiple times
    merged = {}
    for h in holdings:
        sym = h["symbol"]
        if sym in merged:
            merged[sym]["shares"] += h["shares"]
            if h["avg_cost"] is not None:
                merged[sym]["avg_cost"] = h["avg_cost"]
        else:
            merged[sym] = h
            
    return list(merged.values())

def parse_portfolio_report_pdf(text: str) -> List[Dict[str, Any]]:
    """Parse specialized Portfolio Risk Assessment Report PDF."""
    # Preprocess split tickers
    text = text.replace("BAJFINAN\nCE", "BAJFINANCE")
    text = text.replace("BHARTIAR\nTL", "BHARTIARTL")
    text = text.replace("SUNPHAR\nMA", "SUNPHARMA")
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    TICKER_MAPPING = {
        # Crypto
        "BTC": {"symbol": "BTC-USD", "asset_type": "crypto"},
        "ETH": {"symbol": "ETH-USD", "asset_type": "crypto"},
        "SOL": {"symbol": "SOL-USD", "asset_type": "crypto"},
        "BNB": {"symbol": "BNB-USD", "asset_type": "crypto"},
        "LINK": {"symbol": "LINK-USD", "asset_type": "crypto"},
        # US Equities
        "AAPL": {"symbol": "AAPL", "asset_type": "stock"},
        "MSFT": {"symbol": "MSFT", "asset_type": "stock"},
        "NVDA": {"symbol": "NVDA", "asset_type": "stock"},
        "GOOGL": {"symbol": "GOOGL", "asset_type": "stock"},
        "JPM": {"symbol": "JPM", "asset_type": "stock"},
        "JNJ": {"symbol": "JNJ", "asset_type": "stock"},
        "TSLA": {"symbol": "TSLA", "asset_type": "stock"},
        "BRK-B": {"symbol": "BRK-B", "asset_type": "stock"},
        "BRK.B": {"symbol": "BRK-B", "asset_type": "stock"},
        # Indian Equities
        "RELIANCE": {"symbol": "RELIANCE.NS", "asset_type": "stock"},
        "HDFCBANK": {"symbol": "HDFCBANK.NS", "asset_type": "stock"},
        "INFY": {"symbol": "INFY.NS", "asset_type": "stock"},
        "TCS": {"symbol": "TCS.NS", "asset_type": "stock"},
        "ICICIBANK": {"symbol": "ICICIBANK.NS", "asset_type": "stock"},
        "BAJFINANCE": {"symbol": "BAJFINANCE.NS", "asset_type": "stock"},
        "BHARTIARTL": {"symbol": "BHARTIARTL.NS", "asset_type": "stock"},
        "SUNPHARMA": {"symbol": "SUNPHARMA.NS", "asset_type": "stock"},
    }
    
    holdings = []
    for i, line in enumerate(lines):
        upper_line = line.upper()
        if upper_line in TICKER_MAPPING:
            ticker_info = TICKER_MAPPING[upper_line]
            symbol = ticker_info["symbol"]
            asset_type = ticker_info["asset_type"]
            
            if i + 7 < len(lines):
                price_str = lines[i + 1]
                weight_str = lines[i + 7]
                
                price_val = clean_numeric(price_str.replace('■', '').replace('$', ''))
                
                weight_val = None
                if weight_str.endswith('%'):
                    weight_val = clean_numeric(weight_str[:-1])
                
                if price_val and weight_val:
                    shares_val = (weight_val / 100.0) * 100000.0 / price_val
                    holdings.append({
                        "symbol": symbol,
                        "shares": round(shares_val, 4),
                        "avg_cost": round(price_val, 4),
                        "asset_type": asset_type
                    })
    return holdings

def parse_pdf(file_content: bytes) -> List[Dict[str, Any]]:
    """Parse PDF file using pypdf to extract pages text."""
    pdf_file = io.BytesIO(file_content)
    try:
        reader = pypdf.PdfReader(pdf_file)
        full_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
        
        combined_text = "\n".join(full_text)
        if "PORTFOLIO RISK ASSESSMENT REPORT" in combined_text:
            logger.info("Detected PORTFOLIO RISK ASSESSMENT REPORT format. Using specialized parser.")
            return parse_portfolio_report_pdf(combined_text)
            
        return extract_holdings_from_text(combined_text)
    except Exception as e:
        logger.error("Failed parsing PDF: %s", e)
        raise ValueError("Invalid PDF format or unreadable text content.") from e

def parse_docx(file_content: bytes) -> List[Dict[str, Any]]:
    """Parse Microsoft Word DOCX file using python-docx."""
    docx_file = io.BytesIO(file_content)
    try:
        doc = docx.Document(docx_file)
        full_text = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Read table cells
        for table in doc.tables:
            for row in table.rows:
                row_cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells_text:
                    full_text.append(" ".join(row_cells_text))
                    
        return extract_holdings_from_text("\n".join(full_text))
    except Exception as e:
        logger.error("Failed parsing DOCX: %s", e)
        raise ValueError("Invalid DOCX format or unreadable Word content.") from e

def parse_portfolio_file(file_content: bytes, filename: str) -> List[Dict[str, Any]]:
    """Parse portfolio files based on the file extension (case-insensitive)."""
    ext = filename.split('.')[-1].lower()
    
    # 1. Enforce file size limit (10MB)
    if len(file_content) > 10 * 1024 * 1024:
        raise ValueError("File size exceeds the maximum limit of 10MB.")

    if ext == 'csv':
        try:
            # Decode CSV bytes as text
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = file_content.decode('latin-1')
            except Exception as e:
                raise ValueError("Could not decode CSV text content.") from e
        return parse_csv(text)
        
    elif ext == 'pdf':
        return parse_pdf(file_content)
        
    elif ext in ['docx', 'doc']:
        return parse_docx(file_content)
        
    else:
        raise ValueError(f"Unsupported file type '.{ext}'. Supported formats: CSV, PDF, DOCX.")
